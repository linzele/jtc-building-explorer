"""Document generator for JTC land sales tender documents with Azure AI Search integration."""
import os
import io
import re
import requests
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions, ContentSettings
from dotenv import load_dotenv

load_dotenv()

STORAGE_CONN_STR = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
STORAGE_CONTAINER = os.getenv("AZURE_STORAGE_CONTAINER", "land-sales-documents")
AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_ENDPOINT")
AZURE_SEARCH_INDEX = os.getenv("AZURE_SEARCH_INDEX")
AZURE_SEARCH_KEY = os.getenv("AZURE_SEARCH_KEY")

# Location name mapping for search queries
LOCATION_SEARCH_TERMS = {
    "tampines": "tampines",
    "tampines north": "tampines",
    "woodlands": "woodlands",
    "woodlands north": "woodlands",
    "sungei kadut": "sungei kadut",
    "kadut": "sungei kadut",
    "tuas": "tuas",
    "tuas south": "tuas",
    "loyang": "loyang",
    "pioneer": "pioneer",
    "changi": "changi",
    "changi north": "changi",
    "senoko": "senoko",
    "kranji": "kranji",
}


def search_reference_document(location_name: str) -> dict:
    """Search Azure AI Search for reference document content for a location."""
    if not all([AZURE_SEARCH_ENDPOINT, AZURE_SEARCH_INDEX, AZURE_SEARCH_KEY]):
        return None
    
    # Normalize location name for search
    search_term = LOCATION_SEARCH_TERMS.get(location_name.lower(), location_name)
    
    url = f"{AZURE_SEARCH_ENDPOINT}/indexes/{AZURE_SEARCH_INDEX}/docs/search?api-version=2023-11-01"
    headers = {"Content-Type": "application/json", "api-key": AZURE_SEARCH_KEY}
    
    # Search for documents matching the location
    data = {
        "search": search_term,
        "top": 5,
        "select": "title,chunk"
    }
    
    try:
        response = requests.post(url, headers=headers, json=data)
        results = response.json()
        
        # Combine all relevant chunks
        combined_content = []
        source_title = None
        for doc in results.get("value", []):
            title = doc.get("title", "").lower()
            # Filter to only include documents that match the location
            if search_term.lower() in title or any(term in title for term in search_term.lower().split()):
                combined_content.append(doc.get("chunk", ""))
                if not source_title:
                    source_title = doc.get("title", "Reference Document")
        
        if combined_content:
            return {
                "location": location_name,
                "content": "\n\n".join(combined_content),
                "source": source_title
            }
    except Exception as e:
        print(f"Search error: {e}")
    
    return None


def extract_requirements_from_content(content: str) -> dict:
    """Extract key requirements from the reference document content."""
    requirements = {
        "land_use": [],
        "building_height": None,
        "parking": [],
        "setback": [],
        "sustainability": [],
        "prohibited_uses": [],
        "other_requirements": [],
        "development_type": None,
    }
    
    if not content:
        return requirements
    
    content_lower = content.lower()
    
    # Extract development type
    if "semiconductor" in content_lower:
        requirements["development_type"] = "Semiconductor and Advanced Manufacturing"
    elif "wafer" in content_lower:
        requirements["development_type"] = "Wafer Fabrication"
    
    # Extract building height
    height_match = re.search(r'(\d+)\s*m\s*(?:shd|height)', content_lower)
    if height_match:
        requirements["building_height"] = f"{height_match.group(1)}m SHD"
    
    # Extract land use requirements
    if "b2" in content_lower:
        requirements["land_use"].append("B2 Zoning for industrial development")
    if "semiconductor" in content_lower or "wafer" in content_lower:
        requirements["land_use"].append("Semiconductor manufacturing, wafer fabrication, R&D and supporting facilities")
    if "electronics" in content_lower:
        requirements["land_use"].append("Electronics manufacturing")
    
    # Extract prohibited uses
    if "warehousing" in content_lower and "not allowed" in content_lower:
        requirements["prohibited_uses"].append("Warehousing and third-party logistics not allowed")
    if "logistics" in content_lower and "not allowed" in content_lower:
        requirements["prohibited_uses"].append("Third-party logistics operations")
    
    # Extract parking requirements
    if "heavy vehicle" in content_lower:
        requirements["parking"].append("Heavy vehicle parking lots required for operational needs")
    if "overnight" in content_lower and "lorry" in content_lower:
        requirements["parking"].append("Overnight lorry parking to be accommodated")
    
    # Extract sustainability requirements
    if "green mark" in content_lower:
        requirements["sustainability"].append("Green Mark certification required")
    if "solar" in content_lower:
        requirements["sustainability"].append("Solar panel installation requirements")
    
    # Extract setback requirements
    if "setback" in content_lower or "buffer" in content_lower:
        requirements["setback"].append("Green buffer/setback as per Control Plan and JTC requirements")
    
    # Extract other requirements from content
    if "tol" in content_lower or "temporary occupation" in content_lower:
        requirements["other_requirements"].append("Temporary Occupation Licence (TOL) required for works on State Land outside the parcel")
    if "nea" in content_lower:
        requirements["other_requirements"].append("Comply with NEA environmental requirements")
    if "ura" in content_lower:
        requirements["other_requirements"].append("Comply with URA development control guidelines")
    if "bca" in content_lower:
        requirements["other_requirements"].append("Comply with BCA building regulations")
    
    return requirements


def generate_land_sales_tender_document(parcel: dict, buyer: str = "Sample Buyer Pte Ltd", purpose: str = None) -> io.BytesIO:
    """Generate a Word document for land sales tender based on Azure AI Search reference documents."""
    
    # Search for reference document
    ref_doc = search_reference_document(parcel.get("name", ""))
    requirements = extract_requirements_from_content(ref_doc.get("content", "") if ref_doc else "")
    
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SALE OF SITE FOR INDUSTRIAL DEVELOPMENT")
    run.bold = True
    run.font.size = Pt(16)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dev_type = requirements.get("development_type") or purpose or "SEMICONDUCTOR"
    run2 = subtitle.add_run(f"({dev_type.upper()})")
    run2.bold = True
    run2.font.size = Pt(14)
    
    location_para = doc.add_paragraph()
    location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = location_para.add_run(f"LAND PARCEL AT {parcel.get('name', 'INDUSTRIAL AREA').upper()}")
    run3.bold = True
    run3.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph(f"Tender Reference: LSTD-{datetime.now().strftime('%Y%m%d%H%M%S')}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading("CONTENTS", level=1)
    toc_items = [
        "PART I: GENERAL",
        "PART II: PLANNING CONCEPT",
        "PART III: SUMMARY OF PLANNING AND URBAN DESIGN REQUIREMENTS",
        "PART IV: PLANNING AND URBAN DESIGN REQUIREMENTS",
        "PART V: OTHER REQUIRED WORKS",
        "PART VI: OTHER REQUIREMENTS",
        "PART VII: TECHNICAL CONDITIONS",
    ]
    for item in toc_items:
        doc.add_paragraph(f"• {item}")
    doc.add_paragraph()
    
    # PART I: GENERAL
    doc.add_heading("PART I: GENERAL", level=1)
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("1.1 Seller: ").bold = True
    p.add_run("JTC Corporation, a statutory board established under the Jurong Town Corporation Act (Chapter 150), having its principal office at 8 Jurong Town Hall Road, The JTC Summit, Singapore 609434. JTC Corporation is Singapore's leading industrial infrastructure developer, providing comprehensive support for businesses in land, space, and industry development.")
    
    p = doc.add_paragraph()
    p.add_run("1.2 Tenderer: ").bold = True
    p.add_run(f"{buyer}, a company duly incorporated and registered in Singapore under the Companies Act (Chapter 50), having its registered office as stated in the Tender Form. The Tenderer represents and warrants that it has the legal capacity and authority to enter into this tender and perform all obligations thereunder.")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("1.3 Purpose of Tender: ").bold = True
    p.add_run("This document sets forth the terms and conditions for the sale of the Land Parcel described herein for industrial development purposes. The successful Tenderer shall develop the Land Parcel in accordance with the planning parameters, urban design guidelines, and technical conditions specified in this document. The development shall contribute to Singapore's economic growth and align with the nation's strategic industrial development objectives.")
    
    doc.add_paragraph()
    
    # Property Details Table
    doc.add_heading("1.3 Land Parcel Details", level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    details = [
        ("Location", parcel.get('name', 'N/A')),
        ("Postal Code", parcel.get('postal_code', 'N/A')),
        ("Land Area", parcel.get('land_area', 'N/A')),
        ("Zoning", parcel.get('zoning', 'N/A')),
        ("Building Type", parcel.get('building_type', 'Industrial')),
    ]
    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    doc.add_paragraph()
    
    # PART II: PLANNING CONCEPT
    doc.add_heading("PART II: PLANNING CONCEPT", level=1)
    doc.add_paragraph()
    
    dev_type_desc = requirements.get("development_type") or "industrial development"
    doc.add_paragraph(f"2.1 Development Vision: The Land Parcel is strategically designated for {dev_type_desc}. This development forms part of Singapore's comprehensive industrial master plan aimed at strengthening the nation's position as a global hub for advanced manufacturing and technology-driven industries. The development shall be executed in strict accordance with the planning parameters and urban design requirements set out in this document, ensuring alignment with national economic objectives and sustainable development principles.")
    
    doc.add_paragraph(f"2.2 Strategic Location: The Land Parcel at {parcel.get('name', 'the designated area')} is situated within a well-established industrial ecosystem, providing excellent connectivity to major transportation networks, utilities infrastructure, and complementary industrial facilities. This strategic positioning enables efficient supply chain operations and access to a skilled workforce.")
    
    if purpose:
        doc.add_paragraph(f"2.3 Intended Use: The Tenderer has indicated the following intended use for the Land Parcel: {purpose}. This use shall be subject to compliance with all applicable zoning regulations, planning guidelines, and regulatory requirements as specified in this document and as may be required by relevant authorities.")
    
    doc.add_paragraph(f"2.4 Development Standards: The Tenderer shall ensure that the development meets or exceeds the minimum standards set by JTC Corporation, Urban Redevelopment Authority (URA), Building and Construction Authority (BCA), and other relevant regulatory bodies. The development should exemplify best practices in industrial facility design, operational efficiency, and environmental sustainability.")
    
    doc.add_paragraph()
    
    # PART III: SUMMARY OF REQUIREMENTS
    doc.add_heading("PART III: SUMMARY OF PLANNING AND URBAN DESIGN REQUIREMENTS", level=1)
    doc.add_paragraph()
    
    # Create summary table
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Requirement"
    header_cells[1].text = "Details"
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    summary_items = []
    
    # Land Use
    if requirements["land_use"]:
        summary_items.append(("Land Use", "; ".join(requirements["land_use"])))
    else:
        summary_items.append(("Land Use", "As per B2 zoning requirements"))
    
    # Building Height
    summary_items.append(("Building Height", requirements.get("building_height") or "As per Control Plan"))
    
    # Parking
    if requirements["parking"]:
        summary_items.append(("Parking", "; ".join(requirements["parking"])))
    else:
        summary_items.append(("Parking", "Sufficient lots for operational needs"))
    
    # Setback
    if requirements["setback"]:
        summary_items.append(("Setback/Buffer", "; ".join(requirements["setback"])))
    
    # Sustainability
    if requirements["sustainability"]:
        summary_items.append(("Sustainability", "; ".join(requirements["sustainability"])))
    else:
        summary_items.append(("Sustainability", "Green Mark Gold certification required"))
    
    for label, value in summary_items:
        row = summary_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # PART IV: DETAILED REQUIREMENTS
    doc.add_heading("PART IV: PLANNING AND URBAN DESIGN REQUIREMENTS", level=1)
    doc.add_paragraph()
    
    doc.add_heading("4.1 Permitted Uses", level=2)
    doc.add_paragraph("The following uses are permitted on the Land Parcel, subject to compliance with all applicable regulations and obtaining necessary approvals from relevant authorities:")
    if requirements["land_use"]:
        for use in requirements["land_use"]:
            doc.add_paragraph(f"• {use}", style='List Bullet')
    else:
        doc.add_paragraph("• Industrial development as per B2 zoning, including manufacturing, processing, assembly, and ancillary activities directly supporting the primary industrial operations", style='List Bullet')
    doc.add_paragraph("The Tenderer may apply for additional or alternative uses subject to approval from JTC Corporation and relevant authorities. Any change of use shall require prior written consent and may be subject to additional conditions.")
    
    if requirements["prohibited_uses"]:
        doc.add_heading("4.2 Prohibited Uses", level=2)
        doc.add_paragraph("The following uses are expressly prohibited on the Land Parcel and shall not be permitted under any circumstances:")
        for use in requirements["prohibited_uses"]:
            doc.add_paragraph(f"• {use}", style='List Bullet')
        doc.add_paragraph("Any breach of these prohibitions shall constitute a material breach of the tender conditions and may result in termination of the sale agreement and forfeiture of deposits paid.")
    
    doc.add_heading("4.3 Building Requirements", level=2)
    doc.add_paragraph("The Tenderer shall comply with the following building requirements in the design and construction of all structures on the Land Parcel:")
    doc.add_paragraph(f"• Maximum Building Height: {requirements.get('building_height') or 'As per Control Plan'}. This height limit applies to all structures including main buildings, ancillary structures, and any rooftop installations.")
    doc.add_paragraph("• All structures, including rooftop fixtures, mechanical equipment, antennas, signage, and architectural features must comply with the specified height restrictions. No part of any structure shall exceed the maximum permissible height.")
    doc.add_paragraph("• Building design shall incorporate modern industrial architecture that reflects Singapore's position as a global business hub while maintaining functional efficiency for industrial operations.")
    doc.add_paragraph("• Façade treatment shall be of high quality and aesthetically pleasing, contributing positively to the overall industrial landscape of the area.")
    
    if requirements["parking"]:
        doc.add_heading("4.4 Parking Requirements", level=2)
        doc.add_paragraph("The Tenderer shall provide adequate parking facilities to meet operational requirements and comply with the following specifications:")
        for req in requirements["parking"]:
            doc.add_paragraph(f"• {req}", style='List Bullet')
        doc.add_paragraph("• Parking provisions shall be designed to minimize traffic congestion and ensure safe vehicular movement within the premises and on adjacent public roads.")
        doc.add_paragraph("• Adequate loading and unloading bays shall be provided to facilitate efficient logistics operations without causing obstruction to traffic flow.")
    
    doc.add_paragraph()
    
    # PART V: OTHER REQUIRED WORKS
    doc.add_heading("PART V: OTHER REQUIRED WORKS", level=1)
    doc.add_paragraph()
    doc.add_paragraph("5.1 Scope of Works: The Tenderer shall be fully responsible for all works necessary to develop the Land Parcel in accordance with the approved development plans. These works include but are not limited to:")
    doc.add_paragraph()
    doc.add_paragraph("(a) Site Preparation Works: Complete site clearance, demolition of any existing structures, removal of debris, earthworks including cut and fill operations, soil stabilization, and ground improvement works as necessary to prepare the site for construction.")
    doc.add_paragraph()
    doc.add_paragraph("(b) Existing Services: Relocation, diversion, or removal of any existing structures, utilities, services, or infrastructure within the Land Parcel. All such works shall be at the Tenderer's cost and shall be coordinated with the relevant utility providers and authorities.")
    doc.add_paragraph()
    doc.add_paragraph("(c) Utility Connections: Connection to all public utilities including electricity supply from SP Group, water supply from PUB, sewerage and drainage connections, gas supply if required, and telecommunications infrastructure. The Tenderer shall liaise directly with the respective utility providers and bear all connection costs.")
    doc.add_paragraph()
    doc.add_paragraph("(d) Internal Infrastructure: Construction of internal roads, access driveways, parking areas, drainage systems, stormwater management facilities, and all other infrastructure necessary for the proper functioning of the development.")
    doc.add_paragraph()
    doc.add_paragraph("(e) Boundary Works: Construction of boundary walls, fencing, gates, and security infrastructure as required. The design and materials shall comply with JTC's guidelines and be compatible with the surrounding industrial environment.")
    if requirements["other_requirements"]:
        for i, req in enumerate(requirements["other_requirements"]):
            doc.add_paragraph()
            doc.add_paragraph(f"({chr(102+i)}) {req}")
    doc.add_paragraph()
    
    # PART VI: OTHER REQUIREMENTS
    doc.add_heading("PART VI: OTHER REQUIREMENTS", level=1)
    doc.add_paragraph()
    doc.add_paragraph("6.1 Regulatory Compliance: The Tenderer shall comply with all applicable laws, regulations, guidelines, and requirements issued by JTC Corporation, Urban Redevelopment Authority (URA), Building and Construction Authority (BCA), National Environment Agency (NEA), Singapore Land Authority (SLA), Land Transport Authority (LTA), Singapore Civil Defence Force (SCDF), and all other relevant government agencies and statutory boards. It is the Tenderer's responsibility to familiarize themselves with and adhere to all current and future regulatory requirements.")
    doc.add_paragraph()
    doc.add_paragraph("6.2 Approvals and Permits: The Tenderer shall obtain all necessary planning permissions, building plan approvals, environmental permits, and any other approvals or licenses required from relevant authorities before commencement of any development works. Construction shall not commence until all requisite approvals have been obtained and copies submitted to JTC Corporation.")
    doc.add_paragraph()
    doc.add_paragraph("6.3 Environmental Requirements: The Tenderer shall implement appropriate environmental management measures during construction and operation of the development. This includes dust and noise control, proper waste management, prevention of water pollution, and compliance with all NEA environmental standards and requirements.")
    doc.add_paragraph()
    doc.add_paragraph("6.4 Insurance: The Tenderer shall maintain adequate insurance coverage including but not limited to contractor's all risks insurance during construction, public liability insurance, and property insurance upon completion. Evidence of such insurance shall be provided to JTC Corporation upon request.")
    doc.add_paragraph()
    
    # PART VII: TECHNICAL CONDITIONS
    doc.add_heading("PART VII: TECHNICAL CONDITIONS", level=1)
    doc.add_paragraph()
    doc.add_paragraph("7.1 Tender Submission Requirements:")
    doc.add_paragraph("(a) Tender Closing: All tenders must be submitted before the closing date and time as specified in the official Tender Notice published by JTC Corporation. Late submissions will not be accepted under any circumstances.")
    doc.add_paragraph("(b) Tender Deposit: A tender deposit equivalent to 5% of the tendered price must accompany the tender submission. The deposit shall be in the form of a banker's guarantee or cashier's order issued by a bank licensed in Singapore, made payable to JTC Corporation.")
    doc.add_paragraph("(c) Tender Validity: The tender shall remain valid and open for acceptance for a period of ninety (90) days from the tender closing date. The Tenderer shall not withdraw or modify the tender during this validity period.")
    doc.add_paragraph("(d) Tender Documents: The Tenderer shall submit all required documents as specified in the Tender Form, including but not limited to company registration documents, financial statements, development proposal, and track record of similar projects.")
    doc.add_paragraph()
    doc.add_paragraph("7.2 Payment Terms and Schedule:")
    doc.add_paragraph("(a) Initial Deposit: 10% of the accepted tender price shall be paid within fourteen (14) days of the Letter of Acceptance. This deposit is non-refundable except in circumstances as provided in the sale agreement.")
    doc.add_paragraph("(b) Second Payment: An additional 20% of the accepted tender price shall be paid within thirty (30) days from the date of the Letter of Acceptance.")
    doc.add_paragraph("(c) Final Payment: The balance of 70% shall be payable upon completion and handover of the Land Parcel, or as otherwise specified in the sale agreement.")
    doc.add_paragraph("(d) Late Payment: Interest shall be charged on any late payments at the prevailing rate determined by JTC Corporation.")
    doc.add_paragraph()
    doc.add_paragraph("7.3 Development Timeline and Milestones:")
    doc.add_paragraph("(a) Commencement of Construction: The Tenderer shall commence physical construction works within twelve (12) months from the date of possession of the Land Parcel. Failure to commence within this period may result in penalties or termination of the sale agreement.")
    doc.add_paragraph("(b) Completion of Development: The entire development shall be completed, with Temporary Occupation Permit (TOP) obtained, within thirty-six (36) months from the date of possession. Extensions may be granted by JTC Corporation in exceptional circumstances and subject to payment of extension charges.")
    doc.add_paragraph("(c) Progress Reports: The Tenderer shall submit quarterly progress reports to JTC Corporation throughout the construction period, detailing the status of works, any delays encountered, and remedial measures being implemented.")
    doc.add_paragraph("(d) Operational Commencement: The Tenderer shall commence business operations within six (6) months from the date of TOP, demonstrating utilization of the development for its intended purpose.")
    doc.add_paragraph()
    
    # Reference Note
    if ref_doc:
        doc.add_paragraph()
        doc.add_paragraph("─" * 40)
        ref_note = doc.add_paragraph()
        ref_note.add_run("Reference Document: ").bold = True
        ref_note.add_run(ref_doc.get("source", "JTC Industrial Land Sales Document"))
        doc.add_paragraph()
        doc.add_paragraph("This tender document has been generated based on the reference documents available in JTC's repository for the specified location. The Tenderer is advised to refer to the complete set of tender documents and seek clarification from JTC Corporation for any queries.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def upload_to_blob(postal_code: str, buyer: str = "Sample Buyer Pte Ltd", purpose: str = None) -> dict:
    """Generate tender document and upload to Azure Blob Storage."""
    if not STORAGE_CONN_STR:
        raise ValueError("Azure Storage connection string not configured")
    
    from app import LAND_PARCELS
    parcel = next((p for p in LAND_PARCELS if p["postal_code"] == postal_code), None)
    if not parcel:
        raise ValueError(f"Parcel {postal_code} not found")
    
    doc_buffer = generate_land_sales_tender_document(parcel, buyer, purpose)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"LSTD_{postal_code}_{timestamp}.docx"
    
    blob_service = BlobServiceClient.from_connection_string(STORAGE_CONN_STR)
    container = blob_service.get_container_client(STORAGE_CONTAINER)
    try:
        container.create_container()
    except:
        pass
    
    blob = container.get_blob_client(filename)
    doc_buffer.seek(0)
    blob.upload_blob(doc_buffer, overwrite=True, content_settings=ContentSettings(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    ))
    
    sas = generate_blob_sas(
        account_name=blob_service.account_name,
        container_name=STORAGE_CONTAINER,
        blob_name=filename,
        account_key=blob_service.credential.account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=24)
    )
    
    return {"sas_url": f"{blob.url}?{sas}", "filename": filename, "expires_in": "24 hours"}


def generate_local(postal_code: str, buyer: str = "Sample Buyer Pte Ltd", purpose: str = None) -> str:
    """Generate tender document and save locally."""
    from app import LAND_PARCELS
    parcel = next((p for p in LAND_PARCELS if p["postal_code"] == postal_code), None)
    if not parcel:
        raise ValueError(f"Parcel {postal_code} not found")
    
    doc_buffer = generate_land_sales_tender_document(parcel, buyer, purpose)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filepath = os.path.join(os.path.dirname(__file__), 'generated_docs', f"LSTD_{postal_code}_{timestamp}.docx")
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    with open(filepath, 'wb') as f:
        f.write(doc_buffer.getvalue())
    
    return filepath
